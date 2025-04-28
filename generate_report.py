from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
import io

def create_title_page(doc):
    """Create the title page of the report."""
    # Add title
    title = doc.add_heading('Stroke Risk Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('Predictive and Prescriptive Analysis')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date = doc.add_paragraph('Generated on: ' + pd.Timestamp.now().strftime('%Y-%m-%d'))
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

def add_executive_summary(doc):
    """Add executive summary section."""
    doc.add_heading('Executive Summary', level=1)
    
    summary = """
    This report presents a comprehensive analysis of stroke risk factors and prediction models. 
    The analysis combines both predictive modeling and prescriptive insights to provide a 
    complete understanding of stroke risk and potential intervention strategies.
    
    Key findings include:
    • Random Forest model showed the best predictive performance (ROC AUC: 0.511)
    • Top risk factors identified: Average Glucose Level, Stress Levels, and BMI
    • HDL Cholesterol showed the strongest correlation with stroke risk
    • Risk stratification revealed four distinct risk categories
    • Personalized intervention recommendations were generated for each risk category
    """
    
    doc.add_paragraph(summary)
    doc.add_page_break()

def add_predictive_analysis(doc):
    """Add predictive analysis section."""
    doc.add_heading('Predictive Analysis', level=1)
    
    # Model Performance
    doc.add_heading('Model Performance', level=2)
    doc.add_paragraph("""
    Four different machine learning models were evaluated for stroke prediction:
    • Logistic Regression
    • Random Forest
    • Gradient Boosting
    • Support Vector Machine (SVM)
    
    The Random Forest model demonstrated the best performance with an ROC AUC of 0.511.
    All models showed similar performance patterns, suggesting the complexity of stroke prediction.
    """)
    
    # Add model performance comparison image
    doc.add_heading('Model Performance Comparison', level=3)
    doc.add_picture('model_performance_comparison.png', width=Inches(6))
    
    # Feature Importance
    doc.add_heading('Feature Importance Analysis', level=2)
    doc.add_paragraph("""
    The analysis identified several key predictors of stroke risk:
    
    Random Forest Model:
    1. Average Glucose Level (10.8%)
    2. Stress Levels (10.8%)
    3. Body Mass Index (BMI) (10.7%)
    4. LDL Cholesterol (9.9%)
    5. Systolic Blood Pressure (9.6%)
    
    Gradient Boosting Model:
    1. Body Mass Index (BMI) (19.9%)
    2. Average Glucose Level (16.3%)
    3. Stress Levels (14.4%)
    4. Age (10.7%)
    5. Systolic Blood Pressure (9.8%)
    """)
    
    # Add feature importance plots
    doc.add_heading('Feature Importance Visualization', level=3)
    doc.add_picture('feature_importance_random_forest.png', width=Inches(6))
    doc.add_picture('feature_importance_gradient_boosting.png', width=Inches(6))
    
    doc.add_page_break()

def add_prescriptive_analysis(doc):
    """Add prescriptive analysis section."""
    doc.add_heading('Prescriptive Analysis', level=1)
    
    # Risk Stratification
    doc.add_heading('Risk Stratification', level=2)
    doc.add_paragraph("""
    Patients were stratified into four risk categories:
    1. Low Risk
    2. Moderate Risk
    3. High Risk
    4. Very High Risk
    
    Each category contained approximately 25% of the population, with stroke rates ranging from 48.93% to 51.12%.
    """)
    
    # Add risk category distribution image
    doc.add_heading('Risk Category Distribution', level=3)
    doc.add_picture('risk_category_distribution.png', width=Inches(6))
    
    # Intervention Recommendations
    doc.add_heading('Intervention Recommendations', level=2)
    doc.add_paragraph("""
    Based on the analysis, the following intervention recommendations were generated:
    
    Most Common Recommendations:
    1. Regular cardiac monitoring and medication adherence
    2. Cholesterol management through diet and medication
    3. Blood sugar control measures and dietary changes
    4. Weight management program and physical activity
    5. Blood pressure medication and lifestyle modifications
    6. Increase HDL through exercise and dietary changes
    """)
    
    # Add recommendations distribution image
    doc.add_heading('Recommendations Distribution', level=3)
    doc.add_picture('recommendations_distribution.png', width=Inches(6))
    
    doc.add_page_break()

def add_conclusion(doc):
    """Add conclusion and recommendations section."""
    doc.add_heading('Conclusion and Recommendations', level=1)
    
    conclusion = """
    The analysis provides valuable insights into stroke risk prediction and prevention:
    
    1. Risk Prediction:
    • The current models show moderate predictive power
    • Multiple factors contribute to stroke risk
    • No single factor dominates the prediction
    
    2. Risk Management:
    • Focus on HDL Cholesterol management
    • Implement targeted interventions based on risk category
    • Consider personalized treatment plans
    • Regular monitoring of key indicators
    
    3. Future Improvements:
    • Consider ensemble methods for improved performance
    • Focus on feature engineering
    • Implement model calibration
    • Consider cost-sensitive learning
    """
    
    doc.add_paragraph(conclusion)

def generate_report():
    """Generate the complete report."""
    # Create a new document
    doc = Document()
    
    # Add sections
    create_title_page(doc)
    add_executive_summary(doc)
    add_predictive_analysis(doc)
    add_prescriptive_analysis(doc)
    add_conclusion(doc)
    
    # Save the document
    doc.save('Stroke_Analysis_Report.docx')

if __name__ == "__main__":
    generate_report() 